"""Debug script to check actual document structure."""
import sys
import io

sys.path.insert(0, '/Users/life/project/cursor/word-analysis/python-doc-analysis/src')

from docx import Document

def inspect_document():
    """Inspect actual document structure."""
    print("Inspecting real_test.docx structure...")
    
    with open('real_test.docx', 'rb') as f:
        doc = Document(f)
    
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    print("\n" + "="*80)
    print("First 50 paragraphs:")
    print("="*80)
    
    for i, para in enumerate(doc.paragraphs[:50]):
        text = para.text.strip()
        if text:
            print(f"\n[{i}] Text: {text[:100]}...")
            
            # Check for numbering
            element = para._element
            if element.pPr is not None and element.pPr.numPr is not None:
                numId = element.pPr.numPr.numId
                ilvl = element.pPr.numPr.ilvl
                print(f"    Numbering: numId={numId.val if numId else 'None'}, ilvl={ilvl.val if ilvl else 'None'}")
            
            # Check style
            if para.style:
                print(f"    Style: {para.style.name}")
    
    print("\n" + "="*80)
    print("Checking numbering.xml:")
    print("="*80)
    
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is not None:
            xml = numbering_part.element
            print(f"Found numbering.xml with {len(xml.abstractNum)} abstractNum definitions")
            print(f"Found {len(xml.num)} num definitions")
        else:
            print("No numbering.xml found in document")
    except Exception as e:
        print(f"Error accessing numbering: {e}")

if __name__ == "__main__":
    inspect_document()
