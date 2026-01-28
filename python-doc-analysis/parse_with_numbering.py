"""Test script to parse document with heading levels and numbered paths."""
import sys
import os
import json

sys.path.insert(0, '/Users/life/project/cursor/word-analysis/python-doc-analysis/src')

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from doc_analysis.parser.renderer import (
    table_to_html,
    table_to_json,
)

def get_heading_level(para):
    """Determine heading level from paragraph style."""
    if para.style is None:
        return None
    
    style_name = para.style.name.lower() if para.style.name else ""
    
    if "heading 1" in style_name:
        return 1
    elif "heading 2" in style_name:
        return 2
    elif "heading 3" in style_name:
        return 3
    elif "heading 4" in style_name:
        return 4
    elif "heading 5" in style_name:
        return 5
    else:
        return None

def get_heading_label(level):
    """Get heading label like 'h1', 'h2', etc."""
    return f"h{level}"

def parse_document():
    """Parse document using heading styles with numbered paths."""
    print("Parsing real_test.docx with numbered paths...")
    
    with open('real_test.docx', 'rb') as f:
        doc = Document(f)
    
    # Counters for each level
    counters = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    
    result = {
        "document_info": {
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables)
        },
        "sections": []
    }
    
    current_context = {}
    current_section = None
    
    for i, element in enumerate(doc.element.body):
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            text = para.text.strip()
            
            if not text:
                continue
            
            heading_level = get_heading_level(para)
            
            if heading_level is not None:
                # Increment counter for this level
                counters[heading_level] += 1
                
                # Reset deeper level counters
                for level in range(heading_level + 1, 6):
                    counters[level] = 0
                
                # Build number path
                number_parts = []
                for level in range(1, heading_level + 1):
                    number_parts.append(str(counters[level]))
                number_path = ".".join(number_parts)
                
                heading_label = get_heading_label(heading_level)
                
                section_data = {
                    "heading": heading_label,
                    "number_path": number_path,
                    "level": heading_level,
                    "title": text,
                    "parent": None,
                    "content": {
                        "text": [],
                        "tables": [],
                        "images": []
                    }
                }
                
                # Set parent context
                if heading_level > 1:
                    parent_level = heading_level - 1
                    parent_label = get_heading_label(parent_level)
                    if parent_label in current_context:
                        parent_info = current_context[parent_label]
                        section_data["parent"] = {
                            "heading": parent_label,
                            "number_path": parent_info["number_path"],
                            "title": parent_info["title"]
                        }
                
                # Update context
                current_context[heading_label] = {
                    "number_path": number_path,
                    "title": text
                }
                
                # Clear deeper levels
                for l in range(heading_level + 1, 6):
                    label = get_heading_label(l)
                    if label in current_context:
                        del current_context[label]
                
                current_section = section_data
                result["sections"].append(section_data)
                
            elif current_section is not None:
                # Add non-heading content to current section
                current_section["content"]["text"].append(text)
                
        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            
            rows_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                rows_data.append(row_data)
            
            if rows_data and current_section is not None:
                table_data = {
                    "rows": len(rows_data),
                    "cols": len(rows_data[0]) if rows_data else 0,
                    "html": table_to_html(rows_data),
                    "json": table_to_json(rows_data)
                }
                current_section["content"]["tables"].append(table_data)
    
    return result

def main():
    """Main function."""
    json_result = parse_document()
    
    output_file = 'parsed_result_with_numbering.json'
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(json_result, f, ensure_ascii=False, indent=2)
    
    print(f"\nJSON file saved to: {output_file}")
    print(f"Total sections found: {len(json_result['sections'])}")
    
    print("\n" + "="*60)
    print("Preview (first 15 sections):")
    print("="*60)
    
    preview = json_result.copy()
    preview["sections"] = json_result["sections"][:15]
    print(json.dumps(preview, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()
