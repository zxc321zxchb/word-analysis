"""Simple test script for document parser without MySQL."""
import sys
import io
from docx import Document
from docx.shared import Inches

# Add src to path
sys.path.insert(0, '/Users/life/project/cursor/word-analysis/python-doc-analysis/src')

from doc_analysis.parser.docx import DocxParser


def create_test_document():
    """Create a simple test Word document with numbered sections."""
    doc = Document()

    # Title
    doc.add_heading('Test Document', 0)

    # Section 1
    p1 = doc.add_paragraph('1 Introduction')
    p1.style = 'Heading 2'
    doc.add_paragraph('This is the introduction paragraph.')

    # Section 1.1
    p2 = doc.add_paragraph('1.1 Background')
    p2.style = 'Heading 3'
    doc.add_paragraph('Background information here.')

    # Section 2
    p3 = doc.add_paragraph('2 Main Content')
    p3.style = 'Heading 2'

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Header 1'
    table.rows[0].cells[1].text = 'Header 2'
    table.rows[1].cells[0].text = 'Data 1'
    table.rows[1].cells[1].text = 'Data 2'
    table.rows[2].cells[0].text = 'Data 3'
    table.rows[2].cells[1].text = 'Data 4'

    doc.add_paragraph('More content in section 2.')

    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()


def test_parser():
    """Test the document parser."""
    print("Creating test document...")
    doc_content = create_test_document()

    print("\nParsing document...")
    parser = DocxParser()
    parsed = parser.parse(doc_content, "test.docx")

    print(f"\nDocument Info:")
    print(f"  Filename: {parsed.original_filename}")
    print(f"  File Size: {parsed.file_size} bytes")
    print(f"  File Hash: {parsed.file_hash}")
    print(f"  Sections Found: {len(parsed.sections)}")

    print(f"\nSections:")
    for section in parsed.sections:
        print(f"  [{section.number_path}] Level {section.level}: {section.title}")
        print(f"    - Parent: {section.parent_path}")
        print(f"    - Paragraphs: {len(section.paragraphs)}")
        print(f"    - Tables: {len(section.tables)}")
        for p in section.paragraphs:
            print(f"      > {p[:50]}...")
        for table in section.tables:
            print(f"      > Table: {table.rows}x{table.cols}")

    print("\nRendering HTML content:")
    for section in parsed.sections:
        html = section.get_content_html()
        print(f"  [{section.number_path}] HTML: {html[:100]}...")

    print("\nTest completed successfully!")
    return True


if __name__ == "__main__":
    try:
        test_parser()
    except Exception as e:
        print(f"\nTest failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
