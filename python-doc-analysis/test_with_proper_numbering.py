"""Test script to create and parse document with proper numbering."""
import sys
import io
from docx import Document

# Add src to path
sys.path.insert(0, '/Users/life/project/cursor/word-analysis/python-doc-analysis/src')

from doc_analysis.parser.docx import DocxParser

def create_test_document_with_proper_numbering():
    """Create a test Word document with proper numbered sections."""
    doc = Document()

    # Title
    doc.add_heading('Test Document', 0)

    # Section 1 - with proper numbering
    # First, let's create a paragraph with numbering
    # We'll use the built-in list styles
    p1 = doc.add_paragraph('Introduction', style='List Number')
    doc.add_paragraph('This is the introduction paragraph.')

    # Section 1.1 - nested numbering
    p2 = doc.add_paragraph('Background', style='List Number 2')
    doc.add_paragraph('Background information here.')

    # Section 1.2
    p3 = doc.add_paragraph('Objectives', style='List Number 2')
    doc.add_paragraph('Project objectives here.')

    # Section 2
    p4 = doc.add_paragraph('Main Content', style='List Number')
    
    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Header 1'
    table.rows[0].cells[1].text = 'Header 2'
    table.rows[1].cells[0].text = 'Data 1'
    table.rows[1].cells[1].text = 'Data 2'
    table.rows[2].cells[0].text = 'Data 3'
    table.rows[2].cells[1].text = 'Data 4'

    doc.add_paragraph('More content in section 2.')

    # Section 3
    p5 = doc.add_paragraph('Conclusion', style='List Number')
    doc.add_paragraph('This is the conclusion.')

    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()

def test_parser_with_proper_numbering():
    """Test the document parser with proper numbering."""
    print("Creating test document with proper numbering...")
    doc_content = create_test_document_with_proper_numbering()

    print("\nParsing document...")
    parser = DocxParser()
    parsed = parser.parse(doc_content, "test_with_numbering.docx")

    print(f"\nDocument Info:")
    print(f"  Filename: {parsed.original_filename}")
    print(f"  File Size: {parsed.file_size} bytes")
    print(f"  File Hash: {parsed.file_hash}")
    print(f"  Sections Found: {len(parsed.sections)}")

    print(f"\nSections:")
    for section in parsed.sections:
        print(f"  [{section.number_path}] Level {section.level}: {section.title}")
        print(f"    - Parent: {section.parent_path}")
        print(f"    - Paragraphs: {len(section.paragraphs)}")
        print(f"    - Tables: {len(section.tables)}")
        for i, p in enumerate(section.paragraphs):
            print(f"      > Paragraph {i+1}: {p[:100]}...")
        for i, table in enumerate(section.tables):
            print(f"      > Table {i+1}: {table.rows}x{table.cols}")
            print(f"        HTML: {table.html[:100]}...")

    print("\nTest completed successfully!")
    return True

if __name__ == "__main__":
    try:
        test_parser_with_proper_numbering()
    except Exception as e:
        print(f"\nTest failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
