"""Pytest configuration and fixtures."""
import os
import sys
import pytest
from pathlib import Path
from unittest.mock import Mock, MagicMock, patch
from io import BytesIO

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

# Mock init_db BEFORE importing main to prevent startup event from using production DB
with patch("src.doc_analysis.db.session.Base.metadata.create_all"):
    from src.doc_analysis.main import app

from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from fastapi.testclient import TestClient

from src.doc_analysis.db.models import Base
from src.doc_analysis.config import settings


# =============================================================================
# Database Fixtures
# =============================================================================

@pytest.fixture(scope="function")
def db_engine():
    """Create a file-based SQLite database for testing.

    Using file-based database instead of :memory: to ensure
    all connections see the same data in multi-threaded TestClient.
    """
    import tempfile
    import os

    # Create a temporary file for the database
    fd, db_path = tempfile.mkstemp(suffix=".db")
    os.close(fd)

    engine = create_engine(
        f"sqlite:///{db_path}",
        connect_args={"check_same_thread": False},
    )
    Base.metadata.create_all(bind=engine)
    yield engine
    Base.metadata.drop_all(bind=engine)

    # Clean up temporary file
    try:
        os.unlink(db_path)
    except OSError:
        pass


@pytest.fixture(scope="function")
def db_session(db_engine):
    """Create a database session for testing."""
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=db_engine)
    session = TestingSessionLocal()
    yield session
    session.close()


@pytest.fixture(scope="function")
def mock_db_session():
    """Create a mock database session."""
    session = MagicMock()
    session.add = MagicMock()
    session.commit = MagicMock()
    session.refresh = MagicMock()
    session.query = MagicMock()
    session.rollback = MagicMock()
    return session


# =============================================================================
# FastAPI Test Client Fixtures
# =============================================================================

@pytest.fixture(scope="function")
def client(db_engine):
    """Create a test client for FastAPI app."""
    # Override with test session
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=db_engine)

    def get_test_db():
        session = TestingSessionLocal()
        try:
            yield session
        finally:
            session.close()

    from src.doc_analysis.api import routes
    app.dependency_overrides[routes.get_db] = get_test_db

    test_client = TestClient(app)
    yield test_client

    app.dependency_overrides.clear()


@pytest.fixture(scope="function")
def client_no_db():
    """Create a test client without database override."""
    test_client = TestClient(app)
    yield test_client


# =============================================================================
# Document Parser Fixtures
# =============================================================================

@pytest.fixture
def sample_docx_content():
    """Create a sample DOCX file content for testing."""
    from docx import Document

    doc = Document()

    # Add title
    doc.add_heading("Test Document", level=0)

    # Add Section 1
    h1 = doc.add_paragraph("1.0 Introduction")
    h1.style = "Heading 1"
    doc.add_paragraph("This is the introduction.")

    # Add Section 1.1
    h2 = doc.add_paragraph("1.1 Background")
    h2.style = "Heading 2"
    doc.add_paragraph("Background information.")

    # Add Section 2
    h3 = doc.add_paragraph("2.0 Methods")
    h3.style = "Heading 1"
    doc.add_paragraph("Methods description.")

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_large_docx():
    """Create a large DOCX file for testing size limits."""
    from docx import Document

    doc = Document()
    doc.add_heading("Large Document", level=0)

    # Add 100 sections
    for i in range(100):
        h = doc.add_paragraph(f"{i+1}. Section {i+1}")
        h.style = "Heading 1"
        doc.add_paragraph(f"Content for section {i+1}. " * 50)  # Large content

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def invalid_file_content():
    """Create invalid file content for testing error handling."""
    return b"This is not a valid DOCX file."


# =============================================================================
# Mock Fixtures
# =============================================================================

@pytest.fixture
def mock_logger():
    """Create a mock logger."""
    import logging
    from unittest.mock import patch

    with patch("src.doc_analysis.logger.get_logger") as mock:
        logger = MagicMock()
        mock.return_value = logger
        yield logger


@pytest.fixture
def mock_settings():
    """Create mock settings with test values."""
    test_settings = Mock()
    test_settings.database_url = "sqlite:///:memory:"
    test_settings.api_host = "localhost"
    test_settings.api_port = 8000
    test_settings.api_prefix = "/api/v1"
    test_settings.cors_origins = ["http://localhost:3000"]
    test_settings.log_level = "DEBUG"
    test_settings.max_file_size_mb = 50
    test_settings.allowed_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]
    test_settings.max_page_size = 100
    test_settings.default_page_size = 10
    test_settings.max_heading_level = 9
    return test_settings


# =============================================================================
# Test Data Fixtures
# =============================================================================

@pytest.fixture
def sample_document_data():
    """Sample document data for testing."""
    return {
        "id": 1,
        "filename": "test_document.docx",
        "original_filename": "test_document.docx",
        "file_size": 12345,
        "file_hash": "abc123def456",
    }


@pytest.fixture
def sample_section_data():
    """Sample section data for testing."""
    return {
        "id": 1,
        "document_id": 1,
        "number_path": "1.0",
        "level": 0,
        "parent_id": None,
        "title": "Introduction",
        "content_html": "<p>This is the introduction.</p>",
        "content_json": '{"text": "This is the introduction."}',
        "sort_order": 1,
    }
