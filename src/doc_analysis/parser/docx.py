"""Word document parser - main parser module."""
import base64
import io
import hashlib
import re
import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from src.doc_analysis.parser.numbering import (
    WordNumberingExtractor,
)
from src.doc_analysis.logger import get_logger
from src.doc_analysis.config import settings
from src.doc_analysis.parser.renderer import (
    RichTextRenderer,
    RenderedImage,
    RenderedTable,
    table_to_html,
    table_to_json,
    encode_image_to_base64,
)


@dataclass
class ContentItem:
    """A content item with type for ordered rendering."""
    type: str  # "paragraph", "table", "image"
    data: Any


@dataclass
class ParsedSection:
    """A parsed numbered section."""

    heading: str = "h1"
    number_path: str = ""
    level: int = 1
    parent: Optional[Dict[str, str]] = None
    title: Optional[str] = None
    paragraphs: List[str] = field(default_factory=list)          # 纯文本（现有）
    marked_paragraphs: List[str] = field(default_factory=list)   # 带列表标记的新字段
    tables: List[RenderedTable] = field(default_factory=list)
    images: List[RenderedImage] = field(default_factory=list)
    content_items: List[ContentItem] = field(default_factory=list)

    def get_content_html(self) -> str:
        """Get HTML content."""
        renderer = RichTextRenderer()
        for p in self.paragraphs:
            if p.strip():
                renderer.add_paragraph(p)
        return renderer.render_html()

    def get_content_json(self) -> Dict[str, Any]:
        """Get JSON content."""
        renderer = RichTextRenderer()
        for p in self.paragraphs:
            if p.strip():
                renderer.add_paragraph(p)
        return renderer.render_json()


@dataclass
class ParsedDocument:
    """A parsed Word document."""

    original_filename: str
    file_hash: str
    file_size: int
    sections: List[ParsedSection] = field(default_factory=list)


class DocxParser:
    """Parser for Word (.docx) documents."""

    def __init__(self):
        self.logger = get_logger(__name__)
        self.numbering_extractor: Optional[WordNumberingExtractor] = None
        self.sections: List[ParsedSection] = []
        self.current_section: Optional[ParsedSection] = None
        self.section_stack: List[Tuple[str, ParsedSection]] = []
        self._images: List[RenderedImage] = []

    def _get_heading_level(self, paragraph: Paragraph) -> Optional[int]:
        """Determine heading level from paragraph style."""
        if paragraph.style is None:
            return None
        
        style_name = paragraph.style.name.lower() if paragraph.style.name else ""
        
        # Support heading levels 1 to max_heading_level
        for level in range(1, settings.max_heading_level + 1):
            if f"heading {level}" in style_name:
                return level
        return None

    def _get_heading_label(self, level: int) -> str:
        """Get heading label like 'h1', 'h2', etc."""
        return f"h{level}"

    def _get_list_info(self, paragraph: Paragraph) -> Optional[Dict[str, Any]]:
        """检测段落是否为列表项。

        Returns:
            None - 不是列表
            {"type": "ordered", "level": int, "num_id": int} - 有序列表
            {"type": "bullet", "level": int, "num_id": int} - 无序列表
        """
        try:
            element = paragraph._element
            numPr = None

            # First check paragraph's pPr for numPr (explicit numbering)
            if element.pPr is not None and element.pPr.numPr is not None:
                numPr = element.pPr.numPr

            # If not found, check the style's pPr (style-based lists)
            if numPr is None and paragraph.style is not None:
                style_element = paragraph.style._element
                if style_element.pPr is not None and style_element.pPr.numPr is not None:
                    numPr = style_element.pPr.numPr

            if numPr is None:
                return None

            ilvl = numPr.ilvl.val if numPr.ilvl is not None else 0
            num_id = numPr.numId.val if numPr.numId is not None else 0

            # 通过 numbering.xml 判断列表类型
            list_type = self._get_list_type_from_num_id(num_id)

            return {"type": list_type, "level": ilvl, "num_id": num_id}

        except (AttributeError, TypeError):
            return None

    def _get_list_type_from_num_id(self, num_id: int) -> str:
        """根据 num_id 判断列表类型（有序或无序）。"""
        # 如果没有 numbering_extractor，尝试通过默认规则判断
        if self.numbering_extractor is None:
            # 默认的 num_id 规则：
            # - 1-10 通常是有序列表
            # - 11+ 通常是无序列表（子弹列表）
            # 但这只是一个启发式规则
            return "ordered" if num_id <= 10 else "bullet"

        try:
            # 通过 numbering extractor 获取列表格式
            level_def = self.numbering_extractor._get_level_def(num_id, 0)
            if level_def:
                # 检查格式类型
                if level_def.fmt in ("decimal", "roman", "upperRoman", "lowerRoman",
                                     "upperLetter", "lowerLetter", "chineseCounting"):
                    return "ordered"
                elif level_def.fmt in ("bullet", "none"):
                    return "bullet"
                else:
                    # 通过 lvl_text 判断
                    if level_def.lvl_text and "%" in level_def.lvl_text:
                        return "ordered"
                    return "bullet"
        except (AttributeError, TypeError):
            pass

        return "ordered"  # 默认假设为有序列表

    def _format_list_item(self, list_type: str, level: int, counter: int, text: str) -> str:
        """格式化列表项文本。

        Args:
            list_type: "ordered" 或 "bullet"
            level: 缩进级别 (0, 1, 2...)
            counter: 有序列表的序号
            text: 原始文本

        Returns:
            格式化后的文本，如 "<1> 列表内容" 或 "- 列表内容"
        """
        # 计算缩进（每级2个空格）
        indent = "  " * level

        if list_type == "ordered":
            return f"{indent}<{counter}> {text}"
        else:
            return f"{indent}- {text}"

    def parse_by_heading(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse a Word document using heading styles (h1, h2, h3, h4, h5).
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            
        Returns:
            ParsedDocument with all sections
        """
        # Calculate file info
        file_hash = hashlib.sha256(file_content).hexdigest()
        file_size = len(file_content)
        
        # Reset state
        self.sections = []
        self.current_section = None
        self.section_stack = []
        self._images = []
        
        # Counters for each level (support levels 1 to max_heading_level)
        max_level = settings.max_heading_level + 1
        counters = {level: 0 for level in range(1, max_level)}
        current_context = {}
        
        # Load document
        doc = Document(io.BytesIO(file_content))

        # Initialize numbering extractor for list type detection
        self.numbering_extractor = WordNumberingExtractor(doc)

        # List state tracking
        list_state = {
            "num_id": None,      # 当前列表 ID
            "list_type": None,   # "ordered" 或 "bullet"
            "level": 0,          # 当前列表级别
            "counters": {}       # 各级别计数器 {level: count}
        }

        # Extract all images first
        self._extract_all_images(doc)
        # Track which images have been assigned
        assigned_images = set()
        # Reset relationship to image mapping
        self._rel_to_image = {}

        # Iterate through document elements
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Create Paragraph object from element
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()

                # Check if this paragraph contains images
                has_images = False
                for child in paragraph._element.iter():
                    if child.tag.endswith('drawing'):
                        has_images = True
                        break

                # Skip only if both text and images are empty
                if not text and not has_images:
                    continue

                # Collect images in this paragraph
                paragraph_images = []
                if has_images:
                    # Build a mapping from relationship ID to image (only once)
                    if not self._rel_to_image:
                        for rel in doc.part.rels.values():
                            if "image" in rel.target_ref:
                                image_filename = rel.target_ref.split("/")[-1]
                                for img in self._images:
                                    if img.filename == image_filename:
                                        self._rel_to_image[rel.rId] = img
                                        break


                    # Get blip elements to find image references in this paragraph
                    for child in paragraph._element.iter():
                        if child.tag.endswith('blip'):
                            # Extract image relationship ID
                            embed = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed and embed in self._rel_to_image:
                                img = self._rel_to_image[embed]
                                if img.filename not in assigned_images:
                                    paragraph_images.append(img)
                                    assigned_images.add(img.filename)

                heading_level = self._get_heading_level(paragraph)
                
                if heading_level is not None:
                    # Increment counter for this level
                    counters[heading_level] += 1

                    # Reset deeper level counters
                    for level in range(heading_level + 1, max_level):
                        counters[level] = 0
                    
                    # Build number path
                    number_parts = []
                    for level in range(1, heading_level + 1):
                        number_parts.append(str(counters[level]))
                    number_path = ".".join(number_parts)
                    
                    heading_label = self._get_heading_label(heading_level)
                    
                    # Create new section
                    section = ParsedSection(
                        heading=heading_label,
                        number_path=number_path,
                        level=heading_level,
                        title=text,
                    )
                    
                    # Set parent context
                    if heading_level > 1:
                        parent_level = heading_level - 1
                        parent_label = self._get_heading_label(parent_level)
                        if parent_label in current_context:
                            parent_info = current_context[parent_label]
                            section.parent = {
                                "heading": parent_label,
                                "number_path": parent_info["number_path"],
                                "title": parent_info["title"]
                            }
                    
                    # Update context
                    current_context[heading_label] = {
                        "number_path": number_path,
                        "title": text
                    }
                    
                    # Clear deeper levels
                    for l in range(heading_level + 1, max_level):
                        label = self._get_heading_label(l)
                        if label in current_context:
                            del current_context[label]
                    
                    # Reset list state when starting a new section
                    list_state["num_id"] = None
                    list_state["list_type"] = None
                    list_state["level"] = 0
                    list_state["counters"] = {}

                    # Finalize previous section and start new one
                    self._finalize_current_section()
                    self.current_section = section
                    
                elif self.current_section is not None:
                    # Check if this is a list item
                    list_info = self._get_list_info(paragraph)

                    if list_info is not None and text:
                        # This is a list item
                        num_id = list_info["num_id"]
                        list_type = list_info["type"]
                        level = list_info["level"]

                        # Check if this is a new list or continuation
                        if list_state["num_id"] != num_id or list_state["list_type"] != list_type:
                            # New list started
                            list_state["num_id"] = num_id
                            list_state["list_type"] = list_type
                            list_state["counters"] = {}

                        # Update level
                        list_state["level"] = level

                        # Reset counters for deeper levels
                        for l in range(level + 1, 9):
                            if l in list_state["counters"]:
                                del list_state["counters"][l]

                        # Increment counter for current level
                        if level not in list_state["counters"]:
                            list_state["counters"][level] = 0
                        list_state["counters"][level] += 1

                        # Format the list item text
                        counter = list_state["counters"][level]
                        marked_text = self._format_list_item(list_type, level, counter, text)

                        # Add to both paragraphs (plain text) and marked_paragraphs
                        self.current_section.paragraphs.append(text)
                        self.current_section.marked_paragraphs.append(marked_text)
                        self.current_section.content_items.append(ContentItem(type="paragraph", data=text))

                    elif text:
                        # Not a list item - reset list state
                        list_state["num_id"] = None
                        list_state["list_type"] = None
                        list_state["level"] = 0
                        list_state["counters"] = {}

                        # Regular paragraph - add text to current section
                        self.current_section.paragraphs.append(text)
                        self.current_section.marked_paragraphs.append(text)
                        # Add to ordered content items
                        self.current_section.content_items.append(ContentItem(type="paragraph", data=text))

                    # Add images from this paragraph to current section
                    for image in paragraph_images:
                        if image.filename not in [img.filename for img in self.current_section.images]:
                            self.current_section.images.append(image)
                            # Add to ordered content items
                            self.current_section.content_items.append(ContentItem(type="image", data=image))

            elif isinstance(element, CT_Tbl):
                # Create Table object from element
                table = Table(element, doc)
                rendered_table = self._process_table_and_return(table, element)
                if rendered_table and self.current_section is not None:
                    # Add to ordered content items
                    self.current_section.content_items.append(ContentItem(type="table", data=rendered_table))
        
        # Finalize last section
        self._finalize_current_section()

        return ParsedDocument(
            original_filename=filename,
            file_hash=file_hash,
            file_size=file_size,
            sections=self.sections,
        )

    def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse a Word document from bytes.

        Args:
            file_content: Raw file content as bytes
            filename: Original filename

        Returns:
            ParsedDocument with all sections
        """
        # Calculate file info
        file_hash = hashlib.sha256(file_content).hexdigest()
        file_size = len(file_content)

        # Reset state
        self.sections = []
        self.current_section = None
        self.section_stack = []
        self._images = []

        # Load document
        doc = Document(io.BytesIO(file_content))

        # Initialize numbering extractor with the document
        self.numbering_extractor = WordNumberingExtractor(doc)

        # Extract all images first
        self._extract_all_images(doc)

        # Iterate through document elements
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Create Paragraph object from element
                paragraph = Paragraph(element, doc)
                self._process_paragraph(paragraph, element)
            elif isinstance(element, CT_Tbl):
                # Create Table object from element
                table = Table(element, doc)
                self._process_table(table, element)

        # Finalize last section
        if self.current_section is not None:
            self.sections.append(self.current_section)
            self.current_section = None

        return ParsedDocument(
            original_filename=filename,
            file_hash=file_hash,
            file_size=file_size,
            sections=self.sections,
        )

    def _extract_all_images(self, doc: Document) -> None:
        """Extract all images from the document."""
        image_index = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    base64_str = base64.b64encode(image_data).decode("utf-8")
                    mime_type = rel.target_part.content_type

                    # Try to get dimensions
                    width = None
                    height = None
                    # Note: python-docx doesn't easily expose image dimensions
                    # You could use PIL/Pillow if needed

                    image = RenderedImage(
                        filename=rel.target_ref.split("/")[-1],
                        mime_type=mime_type,
                        base64_data=base64_str,
                        width=width,
                        height=height,
                    )
                    self._images.append(image)
                    image_index += 1
                except (KeyError, ValueError, AttributeError) as e:
                    # Skip images with invalid data
                    self.logger.debug(f"Skipping invalid image: {e}")
                except Exception as e:
                    self.logger.warning(f"Unexpected error extracting image: {e}", exc_info=True)

    def _process_paragraph(self, paragraph: Paragraph, element: CT_P) -> None:
        """Process a paragraph element."""
        text = paragraph.text.strip()

        # Check for numbering using the WordNumberingExtractor
        if self.numbering_extractor is None:
            return

        number_info = self.numbering_extractor.get_number_info(paragraph)

        if number_info is not None:
            # This is a numbered heading
            self._finalize_current_section()

            # Get the number path from the extractor (already formatted like "1.2.3")
            number_path = number_info.number_path

            # Derive level from number_path (count the dots)
            # "1" -> level 0, "1.1" -> level 1, "1.1.1" -> level 2
            level = number_path.count(".")

            # Derive parent_path from number_path
            # "1.2.3" -> parent "1.2", "1.2" -> parent "1", "1" -> parent None
            parent_path = None
            if level > 0:
                parent_path = number_path.rsplit(".", 1)[0]

            # Update section stack based on level
            while self.section_stack and len(self.section_stack) > level:
                self.section_stack.pop()

            # Extract title (remove number prefix if present)
            title = text
            match = re.match(r"^[\d.]+\s+(.+)$", text)
            if match:
                title = match.group(1).strip()

            # Create new section
            self.current_section = ParsedSection(
                number_path=number_path,
                level=level,
                parent_path=parent_path,
                title=title if title else None,
            )
            self.section_stack.append((number_path, self.current_section))

        else:
            # Regular paragraph - add to current section
            if self.current_section is not None and text:
                self.current_section.paragraphs.append(text)

    def _process_table(self, table: Table, element: CT_Tbl) -> None:
        """Process a table element."""
        rendered = self._process_table_and_return(table, element)
        if rendered and self.current_section is not None:
            self.current_section.tables.append(rendered)

    def _process_table_and_return(self, table: Table, element: CT_Tbl) -> Optional[RenderedTable]:
        """Process a table element and return the rendered table."""
        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            rows_data.append(row_data)

        if rows_data:
            rendered = RenderedTable(
                rows=len(rows_data),
                cols=len(rows_data[0]) if rows_data else 0,
                html=table_to_html(rows_data),
                json_data=table_to_json(rows_data),
            )
            return rendered
        return None

    def _finalize_current_section(self) -> None:
        """Finalize the current section and add to sections list."""
        if self.current_section is not None:
            self.sections.append(self.current_section)
        self.current_section = None


def parse_docx_file(file_content: bytes, filename: str) -> ParsedDocument:
    """Convenience function to parse a DOCX file.

    Args:
        file_content: Raw file content as bytes
        filename: Original filename

    Returns:
        ParsedDocument
    """
    parser = DocxParser()
    return parser.parse(file_content, filename)
